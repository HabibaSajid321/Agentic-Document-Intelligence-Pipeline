# source/scripts/generate_proposal.py
import json
from pathlib import Path
from docx import Document

def main():
    # Paths
    SCRIPT_DIR = Path(__file__).parent
    SOURCE_DIR = SCRIPT_DIR.parent
    INPUT_DIR = SOURCE_DIR / "input"
    TEMPLATE_DIR = SOURCE_DIR / "templates"
    OUTPUT_DIR = SOURCE_DIR / "output"
    
    INPUT_FILE = INPUT_DIR / "bid_data.json"
    TEMPLATE_FILE = TEMPLATE_DIR / "proposal_template.docx.docx"
    OUTPUT_FILE = OUTPUT_DIR / "Proposal_Output.docx"
    
    # Ensure output dir exists
    OUTPUT_DIR.mkdir(exist_ok=True)
    
    # Load JSON data
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # Build replacements
    replacements = {
        "[CLIENT_NAME]": data["client_name"],
        "[LOCATION]": data["location"],
        "[DATE]": data["date"],
        "[DAYS_VALID]": data["days_valid"],
        "[DIAGNOSTIC_ROOMS]": "\n".join(f"   {room}" for room in data["diagnostic_rooms"]),
        "[MATERIALS_LIST]": "\n".join(f"   {item}" for item in data["materials_list"])
    }
    
    # Load template
    doc = Document(TEMPLATE_FILE)
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    
    # Replace in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save
    doc.save(OUTPUT_FILE)
    print(f"✅ Proposal saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()